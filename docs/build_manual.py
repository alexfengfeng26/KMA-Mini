from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / 'assets' / 'manual'
DIAGRAMS = ROOT / 'diagrams' / 'rendered'
OUT = ROOT / 'KMA-Mini-使用说明手册.docx'
FONT = r'C:\Windows\Fonts\msyh.ttc'

DIAGRAM_SPECS = [
    ('01-end-to-end-overview', '端到端业务总览', ['用户与管理员', '知识门户', '内容治理', '知识技术治理', '运行与安全运维'], '门户连接已发布内容与 AI 问答；内容与技术治理共同支撑知识服务。'),
    ('02-content-lifecycle', '内容生命周期', ['录入/编辑', '草稿', '审核中', '已发布', '门户可见', '下线/失效'], '草稿和审核中内容不可在门户使用；发布后才能检索与问答。'),
    ('03-ai-knowledge-pipeline', 'AI 知识库链路', ['技术文档/已发布内容', '解析与抽取', '分块与关键词', '向量/全文索引', '混合检索', '引用复核', '回答或拒答'], '证据不足时系统明确拒答，不伪造答案。'),
    ('04-permission-access', '权限访问链路', ['本地账号登录', 'RBAC 角色权限', '组织归属', '知识空间 ACL', '内容/功能授权', '门户与后台操作'], '访问必须同时满足角色权限与知识空间 ACL。'),
    ('05-portal-design-release', '门户设计发布链路', ['设计器编辑', '保存草稿', '真实预览', '保存并送审', '审核通过/驳回', '原子发布/回退'], '真实预览只读取指定版本，不改变当前已发布门户。'),
]

def font(size, bold=False):
    return ImageFont.truetype(FONT, size=size, index=1 if bold else 0)

def render_diagram(stem, title, nodes, note):
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    width, height = 1600, 620
    image = Image.new('RGB', (width, height), '#F8FAFC')
    draw = ImageDraw.Draw(image)
    draw.text((72, 48), title, fill='#15395B', font=font(42, True))
    draw.text((72, 112), note, fill='#47657B', font=font(22))
    y, h, gap = 250, 118, 26
    available = width - 144
    box_w = (available - gap * (len(nodes) - 1)) // len(nodes)
    svg_boxes = []
    for index, node in enumerate(nodes):
        x = 72 + index * (box_w + gap)
        draw.rounded_rectangle((x, y, x + box_w, y + h), radius=20, fill='#E7F2FF', outline='#74A9D7', width=3)
        bbox = draw.multiline_textbbox((0, 0), node, font=font(24, True), align='center')
        tx = x + (box_w - (bbox[2] - bbox[0])) / 2
        ty = y + (h - (bbox[3] - bbox[1])) / 2 - 4
        draw.multiline_text((tx, ty), node, fill='#173A5E', font=font(24, True), align='center')
        svg_boxes.append(f'<rect x="{x}" y="{y}" width="{box_w}" height="{h}" rx="20" fill="#E7F2FF" stroke="#74A9D7" stroke-width="3"/><text x="{x + box_w/2}" y="{y + 69}" text-anchor="middle" font-size="24" font-family="Microsoft YaHei" fill="#173A5E">{escape(node)}</text>')
        if index < len(nodes) - 1:
            ax, bx = x + box_w + 5, x + box_w + gap - 5
            draw.line((ax, y + h / 2, bx, y + h / 2), fill='#2F78B7', width=4)
            draw.polygon([(bx, y + h / 2), (bx - 16, y + h / 2 - 9), (bx - 16, y + h / 2 + 9)], fill='#2F78B7')
    image.save(DIAGRAMS / f'{stem}.png')
    svg = f'''<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}"><rect width="100%" height="100%" fill="#F8FAFC"/><text x="72" y="82" font-size="42" font-family="Microsoft YaHei" font-weight="700" fill="#15395B">{escape(title)}</text><text x="72" y="146" font-size="22" font-family="Microsoft YaHei" fill="#47657B">{escape(note)}</text>{''.join(svg_boxes)}</svg>'''
    (DIAGRAMS / f'{stem}.svg').write_text(svg, encoding='utf-8')

def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
    normal = doc.styles['Normal']
    normal.font.name = 'Microsoft YaHei'; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [('Title', 28, '15395B'), ('Heading 1', 18, '15395B'), ('Heading 2', 14, '246B9B'), ('Heading 3', 12, '246B9B')]:
        style = doc.styles[name]
        style.font.name = 'Microsoft YaHei'; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(16); style.paragraph_format.space_after = Pt(8)
    header = section.header.paragraphs[0]
    header.text = 'KMA Mini 使用说明手册  |  单实例演示版'
    header.style = doc.styles['Normal']
    header.runs[0].font.size = Pt(8.5); header.runs[0].font.color.rgb = RGBColor(82, 105, 123)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('KMA Mini · 仅用于学习、演示和内部验证 · 第 ')
    add_field(footer, 'PAGE')
    footer.add_run(' 页')
    return doc

def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = instruction
    separate = OxmlElement('w:fldChar'); separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t'); text.text = '1'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, text, end])

def add_toc(paragraph):
    run = paragraph.add_run()
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.addnext(field)

def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def add_procedure(doc, prereq, steps, verify, note=''):
    doc.add_paragraph('前置权限', style='Heading 3')
    doc.add_paragraph(prereq)
    doc.add_paragraph('操作步骤', style='Heading 3')
    for step in steps:
        doc.add_paragraph(step, style='List Number')
    doc.add_paragraph('结果验证', style='Heading 3')
    doc.add_paragraph(verify)
    if note:
        doc.add_paragraph('注意事项', style='Heading 3')
        doc.add_paragraph(note)

def add_image(doc, path, caption, width=15.5):
    if path.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width))
        cap = doc.add_paragraph(f'图：{caption}'); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9); cap.runs[0].font.color.rgb = RGBColor(82, 105, 123)

def build_docx():
    for spec in DIAGRAM_SPECS: render_diagram(*spec)
    doc = setup_document()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('KMA Mini\n使用说明手册'); r.font.name = 'Microsoft YaHei'; r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = RGBColor(21, 57, 91)
    sub = doc.add_paragraph('全角色使用、治理与本机运维指南'); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(15); sub.runs[0].font.color.rgb = RGBColor(36, 107, 155)
    doc.add_paragraph('\n适用环境：KMA Mini 单实例演示版\n版本日期：2026-07-28', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.add_heading('目录', level=1)
    toc = doc.add_paragraph('在 Word 中右键目录并选择“更新域”，可刷新页码。')
    add_toc(toc)
    doc.add_heading('1 产品与角色', level=1)
    doc.add_paragraph('KMA Mini 是物理单实例 AI 知识库。内容治理决定哪些资料可信、可发布；知识技术治理负责解析、分块、索引、检索和引用。CMS 的 siteKey 仅表示门户站点，不是租户或数据库隔离边界。')
    table = doc.add_table(rows=1, cols=3); table.style = 'Light Shading Accent 1'
    for cell, text in zip(table.rows[0].cells, ['角色', '主要目标', '常用区域']): cell.text = text
    for row in [('普通门户用户','查找资料并获取带引用回答','资料中心、专题、AI 问答'), ('内容管理员','审核与发布可信内容','内容库、审核中心、分类专题'), ('知识技术管理员','保障知识可检索、可评测','空间、技术文档、数据集、检索调试'), ('系统管理员','管理权限与运行稳定性','用户、角色、模型、任务、审计')]:
        cells = table.add_row().cells
        for cell, text in zip(cells, row): cell.text = text
    add_image(doc, ASSETS/'01-login.png', '登录页')
    doc.add_heading('2 门户使用', level=1)
    doc.add_heading('2.1 资料中心、专题与正文', level=2)
    add_procedure(doc, 'content:read，且满足知识空间 ACL。', ['在资料中心按关键词、专题、内容类型或效力状态筛选。', '点击资料打开正文并查看来源、效力与专题。', '在专题学习查看关联内容。'], '只展示已发布、在线且在访问范围内的内容。')
    doc.add_heading('2.2 AI 问答与引用', level=2)
    add_procedure(doc, 'qa:use，且满足资料空间 ACL。', ['打开 AI 问答并输入问题。', '按需限定空间、分类、专题或单篇文档。', '查看回答状态和引用卡片，点击引用跳转正文。'], '证据充分时得到带引用回答；证据不足时得到明确拒答。', '问答是资料辅助，不替代正式审签或业务决策。')
    add_image(doc, ASSETS/'05-portal-ask.png', '门户 AI 问答')
    doc.add_heading('3 内容治理', level=1)
    doc.add_heading('3.1 内容生命周期', level=2)
    add_procedure(doc, '查看 content:read；审核 content:review；发布 content:publish。', ['在内容库新建或编辑内容并完善来源、效力、摘要和专题。', '提交审核；审核人员通过或驳回。', '发布人员发布；必要时下线或修订。'], '草稿和审核中内容不可在门户使用，已发布内容才可检索和问答。', '保存草稿不等于发布。审核、发布和下线为不同权限。')
    add_image(doc, ASSETS/'02-content-governance.png', '内容治理')
    doc.add_heading('4 知识技术治理', level=1)
    doc.add_heading('4.1 知识空间、文档与数据', level=2)
    add_procedure(doc, '空间查看 space:read；技术文档 document:read；数据集 dataset:read。', ['维护知识空间和 ACL。', '在技术文档观察上传、解析、分块与入库状态。', '在数据集与向量查看索引，并在任务与死信处理异常。'], '用户同时满足 RBAC 与 ACL 后才可读取、检索和问答。')
    doc.add_heading('4.2 检索、问答与评测', level=2)
    add_procedure(doc, 'retrieval:use、qa:use、evaluation:read。', ['在检索调试查看召回与过滤。', '在问答实验室核验回答和引用。', '在 RAG 评测查看 Recall@K、MRR、正确率、引用准确率与拒答率。'], '所有高质量回答均可回溯其资料依据。')
    add_image(doc, ASSETS/'03-knowledge-technology.png', '知识技术治理')
    doc.add_heading('5 门户运营', level=1)
    doc.add_heading('5.1 设计、预览与发布', level=2)
    add_procedure(doc, '编辑需要 portal-site:update 与 portal-page:edit；审核、发布需要对应权限。', ['在门户设计中心编辑结构、属性、主题和内容范围。', '保存草稿后用预览变更打开真实门户版本。', '确认后保存并送审；审核通过后由发布人员原子发布。'], '预览页面显示“预览中 · Vx”，使用真实资料但不改变当前已发布版本。', 'AI 设计只生成候选提案；应用后仍需保存、审核和发布。')
    add_image(doc, ASSETS/'04-portal-designer.png', '门户设计中心')
    doc.add_heading('6 系统管理与运维', level=1)
    doc.add_heading('6.1 用户、角色与组织', level=2)
    add_procedure(doc, '用户、角色、组织分别需要相应管理权限。', ['创建用户并分配最小必要角色。', '维护组织归属。', '通过知识空间 ACL 授予内容范围。'], '路由与 API 都执行授权检查，权限变更刷新会话后生效。')
    doc.add_heading('6.2 本机启动、监测与故障处理', level=2)
    add_bullets(doc, ['数据库目标必须为 kma_mini，并启用 vector 扩展。', '通过进程环境变量提供数据库密码、JWT 密钥和模型密钥，禁止提交到仓库。', '启动 API（8090）、前端（27183）与按需 Worker；Readiness 应返回 UP。', '门户无内容时检查发布状态、ACL、内容范围与效力；问答无引用时检查解析/分块、索引、模型与任务。', '任务积压或死信时检查 Worker、租约、重试次数及外部模型/存储依赖。'])
    doc.add_heading('7 核心流程图', level=1)
    for stem, title, _, _ in DIAGRAM_SPECS:
        doc.add_heading(title, level=2)
        add_image(doc, DIAGRAMS / f'{stem}.png', title, 16.0)
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    build_docx()
